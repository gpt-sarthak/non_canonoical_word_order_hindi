"""
generate_doc.py

Generates a Word document explaining the full replication pipeline
in plain English with logical detail.

Usage:
    python scripts/generate_doc.py

Output:
    reports/pipeline_explanation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "reports/pipeline_explanation.docx"


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Non-Canonical Word Order in Hindi", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Replication of Ranjan & van Schijndel (2024)\nFull Pipeline — Plain English Explanation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)

    doc.add_page_break()

    # ── 1. What is this project? ─────────────────────────────────────────────
    add_heading(doc, "1. What Is This Project?", 1)
    add_body(doc,
        "This project is a replication of a research paper published in 2024 by Ranjan and "
        "van Schijndel, titled \"Does Dependency Locality Predict Non-canonical Word Order in Hindi?\". "
        "The paper asks a simple but deep question: when a Hindi speaker places a word in an unusual "
        "position in a sentence — for example, moving the object to the front — what drives that choice? "
        "Is it because that word order is easier to process? Or is it because of discourse reasons, such "
        "as the speaker already mentioned that object earlier in the conversation?"
    )
    add_body(doc,
        "Hindi is a relatively free word-order language. The default order is Subject-Object-Verb (SOV), "
        "e.g. 'Ram apple ate'. But speakers sometimes front the object: 'Apple Ram ate' (DOSV — direct "
        "object first), or front the indirect object: 'To-Ram apple gave' (IOSV). These non-canonical "
        "orderings are the focus of the paper."
    )
    add_body(doc,
        "The project tests whether several computational measures — how surprising a word order is, "
        "how long the grammatical links between words are, and whether the information flows naturally "
        "from known to new — can predict which word order a speaker will actually use."
    )

    # ── 2. The Data ───────────────────────────────────────────────────────────
    add_heading(doc, "2. The Data", 1)

    add_heading(doc, "2a. Main Evaluation Data — HDTB Treebank", 2)
    add_body(doc,
        "The core dataset is the Hindi Dependency Treebank (HDTB), also known as the Hindi-Urdu "
        "Treebank (HUTB). This is a hand-annotated collection of real Hindi sentences drawn from "
        "newspaper articles. Each sentence has been carefully annotated by linguists to record:"
    )
    add_bullet(doc, "Every word (token) in the sentence")
    add_bullet(doc, "The grammatical role of each word (subject, object, verb, etc.)")
    add_bullet(doc, "Which other word each word depends on (the dependency head)")
    add_bullet(doc, "The part of speech of each word (noun, verb, adjective, etc.)")
    add_body(doc,
        "The treebank is stored in CoNLL-U format — a standard plain-text table format used in "
        "computational linguistics. Each sentence is a block of lines, one per word, with columns "
        "for the word form, lemma, part of speech, grammatical relation, and head."
    )
    add_body(doc,
        "Total raw sentences: 13,306. After applying the paper's filters (see Section 4), "
        "2,828 sentences remain. These 2,828 sentences are used to generate 92,299 sentence pairs "
        "(each real sentence paired with a permuted version of itself)."
    )
    add_body(doc,
        "Important: the HDTB is used ONLY for evaluation — to compute features and test whether "
        "the models can predict which word order is the real one. It is not used to train the "
        "language models."
    )

    add_heading(doc, "2b. Training Data — Hindi Wikipedia", 2)
    add_body(doc,
        "The two language models (Trigram and LSTM) are trained on a large corpus of Hindi text "
        "extracted from Hindi Wikipedia. This gives the models a broad statistical knowledge of "
        "Hindi — which word sequences are common, which words tend to follow which others, and "
        "so on."
    )
    add_body(doc,
        "The Wikipedia text is preprocessed by:"
    )
    add_bullet(doc, "Extracting article text from the Wikipedia XML dump")
    add_bullet(doc, "Splitting on sentence boundary markers (Hindi uses '।' as a full stop, plus '?', '!')")
    add_bullet(doc, "Cleaning: removing XML tags, URLs, brackets, and punctuation")
    add_bullet(doc, "Result: 2,655,055 sentences saved to wiki_sentences.txt")
    add_body(doc,
        "A vocabulary of 30,000 most frequent words is built from these sentences. "
        "Any word not in this vocabulary is replaced with a special <UNK> token during training."
    )

    add_table(doc,
        ["Dataset", "Purpose", "Size"],
        [
            ["HDTB Treebank (hi_hdtb-ud-train.conllu)", "Evaluation — generate sentence pairs and compute features", "13,306 raw → 2,828 valid → 92,299 pairs"],
            ["Hindi Wikipedia (wiki_sentences.txt)", "Train Trigram LM and LSTM LM", "2,655,055 sentences"],
            ["Vocabulary (vocab.pkl)", "Shared word index for LSTM", "30,000 words + <UNK>, <SOS>, <EOS>"],
        ]
    )

    # ── 3. The Pipeline Overview ──────────────────────────────────────────────
    add_heading(doc, "3. Pipeline Overview", 1)
    add_body(doc,
        "The pipeline has three distinct phases:"
    )
    add_bullet(doc, "Phase 1 — Preprocessing: Train the language models on Wikipedia")
    add_bullet(doc, "Phase 2 — Feature Extraction: Use the models to score every sentence pair from the treebank")
    add_bullet(doc, "Phase 3 — Analysis: Train a logistic regression model and compare results to the paper")
    add_body(doc,
        "Each phase is described in detail below."
    )

    # ── 4. Phase 1 — Preprocessing ───────────────────────────────────────────
    add_heading(doc, "4. Phase 1 — Preprocessing (Training the Language Models)", 1)

    add_heading(doc, "4a. Building the Vocabulary", 2)
    add_body(doc,
        "Before training any model, we count how often each word appears across all 2.6 million "
        "Wikipedia sentences. We keep the 30,000 most frequent words. Everything else becomes <UNK>. "
        "This vocabulary is saved as vocab.pkl and reused by both the LSTM and the adaptive model."
    )

    add_heading(doc, "4b. Training the Trigram Language Model", 2)
    add_body(doc,
        "A trigram language model estimates the probability of a word given the two words "
        "immediately before it. For example, given the words 'राम ने', it estimates how likely "
        "each possible next word is. We train this model on all 2.6 million Wikipedia sentences "
        "using NLTK's MLE (Maximum Likelihood Estimation) class."
    )
    add_body(doc,
        "MLE means: the probability of a word following a context is simply the count of times "
        "that trigram appeared, divided by the count of times that context appeared. No smoothing "
        "is applied by the model itself — instead, our scoring code uses a three-level backoff:"
    )
    add_bullet(doc, "Try trigram P(word | prev2, prev1) first")
    add_bullet(doc, "If unseen, fall back to bigram P(word | prev1)")
    add_bullet(doc, "If still unseen, fall back to unigram P(word)")
    add_bullet(doc, "If completely unknown, use a tiny epsilon (1e-12) to avoid log(0)")
    add_body(doc,
        "The trained model is saved as models/trigram/trigram.pkl."
    )

    add_heading(doc, "4c. Training the LSTM Language Model", 2)
    add_body(doc,
        "An LSTM (Long Short-Term Memory) is a type of neural network designed to process "
        "sequences. Unlike the trigram model which only looks at the two previous words, "
        "the LSTM maintains a hidden state that in principle captures longer-range context."
    )
    add_body(doc,
        "Architecture:"
    )
    add_bullet(doc, "Embedding layer: maps each word index to a 256-dimensional vector")
    add_bullet(doc, "LSTM layer: 2 layers, 256 hidden units each, with 30% dropout during training")
    add_bullet(doc, "Output layer: projects the hidden state back to vocabulary size (30,000)")
    add_body(doc,
        "Training:"
    )
    add_bullet(doc, "Trained on 2 million tokens from Wikipedia (a subset of the full corpus)")
    add_bullet(doc, "2 epochs, batch size 128, sequence length 50, Adam optimizer (lr=0.001)")
    add_bullet(doc, "Task: predict the next word at every position (standard language modelling)")
    add_bullet(doc, "Epoch 1 loss: 4.42  →  Epoch 2 loss: 3.36")
    add_body(doc,
        "The trained model weights are saved as models/lstm/base_model.pt. "
        "Checkpointing was added to prevent full restart if training is interrupted — "
        "after each epoch a checkpoint.pt is saved, and on next run the model resumes "
        "from the last completed epoch."
    )

    # ── 5. Phase 2 — Feature Extraction ──────────────────────────────────────
    add_heading(doc, "5. Phase 2 — Feature Extraction", 1)
    add_body(doc,
        "This is the core of the pipeline. For each of the 2,828 valid treebank sentences, "
        "we generate permuted variants and then score both the original sentence and each "
        "variant on five features. The result is a table of 92,299 rows saved as features.csv."
    )

    add_heading(doc, "5a. Sentence Filtering", 2)
    add_body(doc,
        "Not all 13,306 treebank sentences are usable. The paper defines strict criteria for "
        "which sentences can participate in the experiment. A sentence is kept only if:"
    )
    add_bullet(doc, "Its root verb is finite (VerbForm=Fin) — ensures we have a proper main clause")
    add_bullet(doc, "The root is a VERB or AUX — filters out nominal sentences")
    add_bullet(doc, "The root has at least two pre-verbal nominal dependents (subject, object, or indirect object) directly attached to it — needed to generate meaningful permutations")
    add_bullet(doc, "The sentence is projective — no crossing dependency arcs (a linguistic well-formedness condition)")
    add_bullet(doc, "No negative markers (नहीं, न, मत) — negation complicates the analysis")
    add_body(doc,
        "13,306 → 2,828 sentences pass all filters (21.2% retention rate)."
    )

    add_heading(doc, "5b. Variant Generation", 2)
    add_body(doc,
        "For each valid sentence, we generate alternative word orderings by permuting the "
        "pre-verbal phrases while keeping the verb in its original position. The logic is:"
    )
    add_bullet(doc, "Identify all pre-verbal phrases directly dependent on the root verb")
    add_bullet(doc, "Generate all permutations of those phrases")
    add_bullet(doc, "Filter out permutations that produce impossible adjacencies (e.g. a postposition landing next to an incompatible element). This filter is learned from the actual treebank — pairs of adjacent dependency relations that never appear in real data are forbidden.")
    add_bullet(doc, "Remove any variant that is identical to the reference sentence")
    add_bullet(doc, "If more than 99 variants survive, randomly sample exactly 99")
    add_body(doc,
        "Each (reference, variant) pair is then labelled by construction type:"
    )
    add_bullet(doc, "SOV — subject comes before object (canonical Hindi order)")
    add_bullet(doc, "DOSV — direct object appears first (non-canonical)")
    add_bullet(doc, "IOSV — indirect object appears first (non-canonical)")

    add_heading(doc, "5c. Feature 1 — Dependency Length (DL)", 2)
    add_body(doc,
        "Dependency length is the distance between a word and the word it depends on, measured "
        "in number of positions. For example, if 'ate' is at position 5 and depends on 'Ram' at "
        "position 1, the dependency length is |5 - 1| = 4. Shorter dependency lengths are "
        "generally easier to process because the brain doesn't have to hold the incomplete "
        "dependency open for as long."
    )
    add_body(doc,
        "For each sentence, we sum the dependency lengths of all words (excluding the root, "
        "which has no head). delta_dl = DL(reference) − DL(variant). A negative delta means "
        "the reference has shorter total dependency length — which is the theoretically "
        "expected direction if DL drives word order choice."
    )

    add_heading(doc, "5d. Feature 2 — Trigram Surprisal", 2)
    add_body(doc,
        "Surprisal measures how unexpected a word is given its context. A word with low "
        "surprisal was predictable; a word with high surprisal was unexpected. Sentence-level "
        "surprisal is the sum of per-word surprisals: -log P(word | two previous words)."
    )
    add_body(doc,
        "The trigram model trained on Wikipedia provides these probabilities. "
        "delta_trigram = surprisal(reference) − surprisal(variant). A negative delta means "
        "the reference sentence is more predictable — the speaker chose the word order that "
        "flows more naturally according to learned Hindi statistics."
    )

    add_heading(doc, "5e. Feature 3 — LSTM Surprisal", 2)
    add_body(doc,
        "Same idea as trigram surprisal, but using the LSTM language model instead of the "
        "trigram model. The LSTM processes the sentence left to right, maintaining a hidden "
        "state, and at each position predicts the next word. Surprisal = -log P(word | all "
        "previous words in the LSTM's hidden state)."
    )
    add_body(doc,
        "The LSTM is more powerful than the trigram because it can theoretically capture "
        "longer dependencies. delta_lstm = surprisal(reference) − surprisal(variant)."
    )

    add_heading(doc, "5f. Feature 4 — Adaptive LSTM Surprisal", 2)
    add_body(doc,
        "The adaptive model starts from the base LSTM but fine-tunes itself on each "
        "document's context sentences before scoring the target sentence. Concretely:"
    )
    add_bullet(doc, "The base LSTM is copied fresh for each unique reference sentence")
    add_bullet(doc, "The model runs a single gradient update step on the preceding context sentences")
    add_bullet(doc, "This makes the model adapt its expectations to the specific topic and vocabulary of the current document")
    add_bullet(doc, "Then it scores the reference and variant sentences with the adapted weights")
    add_body(doc,
        "The motivation: if a word has been used several times in the preceding discourse, "
        "the adapted model will assign it lower surprisal on the next occurrence — capturing "
        "a form of discourse coherence beyond what the static LSTM knows."
    )

    add_heading(doc, "5g. Feature 5 — Information Status (IS / Givenness)", 2)
    add_body(doc,
        "Information Status captures whether the entities being talked about are already "
        "known to the listener (GIVEN) or newly introduced (NEW). In natural discourse, "
        "languages tend to place given information before new information."
    )
    add_body(doc,
        "For each sentence, we look at the subject and object phrases. A phrase is GIVEN if:"
    )
    add_bullet(doc, "Its head word is a pronoun (pronouns refer back to something already mentioned), OR")
    add_bullet(doc, "At least one of its content words appeared in the immediately preceding sentence")
    add_body(doc,
        "The IS score is then:"
    )
    add_bullet(doc, "+1 if the first (leftmost) phrase is GIVEN and the second is NEW — canonical given-before-new order")
    add_bullet(doc, "-1 if the first phrase is NEW and the second is GIVEN — marked order")
    add_bullet(doc, " 0 if both are GIVEN or both are NEW — no informational contrast")
    add_body(doc,
        "delta_is = IS(reference) − IS(variant). A positive delta means the reference "
        "follows given-before-new ordering more than its variant."
    )

    add_heading(doc, "5h. The features.csv Output", 2)
    add_body(doc,
        "After all five features are computed, everything is saved to data/features/features.csv. "
        "Each row represents one (reference, variant) pair with the following columns:"
    )
    add_table(doc,
        ["Column", "Description"],
        [
            ["sentence_id", "Index of the original treebank sentence"],
            ["construction_type", "SOV / DOSV / IOSV — what word order the reference uses"],
            ["reference", "The actual Hindi sentence from the treebank"],
            ["variant", "The permuted version of the sentence"],
            ["dl_reference / dl_variant", "Total dependency length of each sentence"],
            ["delta_dl", "DL(reference) − DL(variant)"],
            ["trigram_reference / trigram_variant", "Total trigram surprisal of each sentence"],
            ["delta_trigram", "Surprisal(reference) − Surprisal(variant)"],
            ["lstm_reference / lstm_variant", "Total LSTM surprisal of each sentence"],
            ["delta_lstm", "Same for LSTM"],
            ["adaptive_reference / adaptive_variant", "Total adaptive LSTM surprisal"],
            ["delta_adaptive", "Same for adaptive LSTM"],
            ["is_reference / is_variant", "IS score of each sentence"],
            ["delta_is", "IS(reference) − IS(variant)"],
        ]
    )

    # ── 6. Phase 3 — Analysis ─────────────────────────────────────────────────
    add_heading(doc, "6. Phase 3 — Analysis and the Ranking Model", 1)

    add_heading(doc, "6a. The Pairwise Ranking Formulation", 2)
    add_body(doc,
        "The central question is: given a reference sentence and one of its variants, can the "
        "model predict which one is the real sentence (the one actually produced by the speaker)?"
    )
    add_body(doc,
        "This is treated as a pairwise classification problem. For each pair (reference, variant), "
        "the model receives the delta values (one number per feature) and must predict: "
        "'reference preferred' or 'variant preferred'. Because the reference is always the "
        "real sentence, the label is always 'reference preferred' — but to prevent the model "
        "from cheating, each pair is duplicated with flipped signs: the original (delta → label=1) "
        "and its negation (-delta → label=0). This forces the model to learn actual feature "
        "relationships rather than just always predicting one class."
    )

    add_heading(doc, "6b. Logistic Regression", 2)
    add_body(doc,
        "The ranking model is logistic regression — a simple, interpretable model that fits "
        "a weighted combination of the features to predict the probability of 'reference preferred'. "
        "The features are z-scored (standardised) so coefficients are on the same scale and "
        "directly comparable."
    )
    add_body(doc,
        "Accuracy is measured using 10-fold cross-validation: the data is split into 10 equal "
        "chunks, the model is trained on 9 chunks and tested on the 10th, repeated 10 times. "
        "The random split is at the sentence level (not the pair level) to prevent data leakage "
        "between train and test."
    )

    add_heading(doc, "6c. Results vs the Paper", 2)
    add_table(doc,
        ["Condition", "Our Accuracy", "Paper Accuracy", "Gap"],
        [
            ["Full dataset — baseline (adaptive + trigram)", "80.80%", "85.18%", "-4.38%"],
            ["Full dataset — all features", "81.49%", "85.04%", "-3.55%"],
            ["DOSV — baseline", "68.96%", "81.24%", "-12.28%"],
            ["DOSV — all features", "73.60%", "80.46%", "-6.86%"],
            ["IOSV — baseline", "82.28%", "89.43%", "-7.15%"],
            ["IOSV — all features", "82.62%", "90.02%", "-7.40%"],
        ]
    )
    add_body(doc,
        "The main source of the gap is the missing PCFG surprisal feature (see Section 7). "
        "All coefficient directions match the paper perfectly (10 out of 10)."
    )

    add_heading(doc, "6d. What the Coefficients Tell Us", 2)
    add_body(doc,
        "The logistic regression coefficients tell us how strongly each feature predicts "
        "that the reference sentence is preferred. A negative coefficient means the reference "
        "tends to have a lower value of that feature — i.e. the speaker chooses the word order "
        "that minimises that measure. A positive coefficient means the reference tends to have "
        "a higher value."
    )
    add_table(doc,
        ["Feature", "DOSV coeff", "IOSV coeff", "Interpretation"],
        [
            ["Trigram Surprisal", "-0.77", "-0.65", "Reference is more predictable — speaker minimises surprisal"],
            ["Adaptive Surprisal", "-0.27", "-1.11", "Reference fits discourse context better"],
            ["LSTM Surprisal", "+0.27", "+1.11", "Reference is less predictable by LSTM — adaptive absorbs the signal"],
            ["Dependency Length", "-0.81", "-0.29", "Reference has shorter dependencies — stronger for DOSV"],
            ["Information Status", "+0.41", "+0.06", "Reference follows given-before-new — significant for DOSV"],
        ]
    )
    add_body(doc,
        "Key finding: for DOSV (direct object fronting), dependency length is the dominant "
        "predictor — speakers front the direct object when it shortens the dependency links. "
        "For IOSV (indirect object fronting), surprisal dominates — the word order that fits "
        "the discourse context better is chosen. This replicates the paper's main finding."
    )

    # ── 7. Known Gap ──────────────────────────────────────────────────────────
    add_heading(doc, "7. Known Gap — PCFG Surprisal", 1)
    add_body(doc,
        "The paper includes a fifth surprisal predictor: PCFG (Probabilistic Context-Free "
        "Grammar) surprisal. This is computed from a constituency parse tree — a "
        "different type of syntactic representation from the dependency trees we use. "
        "It captures phrase-structure expectations: how likely is it that a sentence has "
        "the bracketing structure it does, given a grammar learned from a treebank?"
    )
    add_body(doc,
        "Our treebank (HDTB) is in dependency format. Converting dependency trees to "
        "constituency trees for Hindi requires a non-trivial conversion step, and a "
        "Hindi PCFG parser is not straightforward to obtain. The paper shows that for "
        "IOSV sentences, PCFG surprisal absorbs most of the dependency length effect — "
        "meaning without PCFG, DL appears significant for both DOSV and IOSV in our "
        "replication, whereas the paper shows it is only significant for DOSV."
    )
    add_body(doc,
        "The pcfg_features.py file exists as a scaffold but currently returns a placeholder "
        "value (log of sentence length). It is not included in the features.csv."
    )

    # ── 8. How to Run ─────────────────────────────────────────────────────────
    add_heading(doc, "8. How to Run the Full Pipeline", 1)
    add_body(doc,
        "All commands are run from the project root with PYTHONPATH=. to ensure Python "
        "can find the project's modules."
    )
    add_table(doc,
        ["Step", "Command", "Output"],
        [
            ["Extract Wikipedia text", "python preprocessing/wiki_extract.py", "data/processed/wiki_sentences.txt"],
            ["Build vocabulary", "python preprocessing/build_vocab.py", "data/processed/vocab.pkl"],
            ["Train trigram LM", "python models/trigram/train_trigram.py", "models/trigram/trigram.pkl"],
            ["Train LSTM LM", "python models/lstm/train_base_model.py", "models/lstm/base_model.pt"],
            ["Build feature dataset", "PYTHONPATH=. python scripts/build_feature_dataset.py", "data/features/features.csv"],
            ["Train ranking model", "PYTHONPATH=. python scripts/train_ranking_model.py", "Printed results"],
            ["Compare with paper", "PYTHONPATH=. python scripts/compare_with_paper.py", "reports/paper_comparison.txt"],
            ["Generate HTML report", "PYTHONPATH=. python scripts/analyse_results.py", "reports/pipeline_analysis.html"],
            ["Run all tests", "PYTHONPATH=. python tests/test_dl.py (etc.)", "tests/output_test/"],
        ]
    )

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
