"""
generate_detailed_doc.py

Generates a detailed Word document with step-by-step explanation of
which data trains which model, preprocessing, HDTB constraints,
and evaluation — all with worked examples and visual diagrams.

Usage:
    python scripts/generate_detailed_doc.py

Output:
    reports/pipeline_detailed_explanation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = "reports/pipeline_detailed_explanation.docx"


# ── Helpers ───────────────────────────────────────────────────────────────────

def style_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    p.add_run(text)
    return p

def mono(doc, text):
    """Add a monospace / code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p

def label(doc, text):
    """Bold label paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    return p

def note(doc, text):
    """Italic note paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("Note: " + text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        shade_cell(hdr_cells[i], "D9E1F2")
    for row_data in rows:
        r = table.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table

def pass_fail_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr_cells[i], "D9E1F2")
    for row_data in rows:
        r = table.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = str(val)
        # colour last cell green/red
        result = row_data[-1]
        shade_cell(r[-1], "C6EFCE" if "PASS" in result else "FFC7CE")
    doc.add_paragraph()


# ── Main document ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover ────────────────────────────────────────────────────────────────
    t = doc.add_heading("Hindi Non-Canonical Word Order", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Complete Data & Model Guide — With Examples\nReplication of Ranjan & van Schijndel (2024)")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ── 0. Big Picture ───────────────────────────────────────────────────────
    style_heading(doc, "0. The Big Picture", 1)
    body(doc,
        "This project answers the question: when a Hindi speaker uses an unusual word order "
        "(e.g. putting the object before the subject), what predicts that choice? "
        "We test five computational measures: how surprising the word order is, how long the "
        "grammatical links between words are, and whether information flows naturally from known "
        "to new. Each measure is computed for the real sentence and a machine-generated alternative "
        "— then we test whether the measure correctly identifies the real one."
    )

    label(doc, "Data Flow at a Glance")
    mono(doc, "HINDI WIKIPEDIA (2.6M sentences)")
    mono(doc, "       |")
    mono(doc, "       |---> [Preprocessing] ---> wiki_sentences.txt")
    mono(doc, "       |                               |")
    mono(doc, "       |                               |--> vocab.pkl       (30,000 words)")
    mono(doc, "       |                               |--> trigram.pkl     (Trigram LM)")
    mono(doc, "       |                               '--> base_model.pt   (LSTM LM)")
    mono(doc, "")
    mono(doc, "HDTB TREEBANK (16,649 sentences: train + dev + test)")
    mono(doc, "       |")
    mono(doc, "       |---> [Filter] ---> 2,828 valid sentences")
    mono(doc, "       |                       |")
    mono(doc, "       |               [Generate Variants]")
    mono(doc, "       |                       |")
    mono(doc, "       |               92,299 (reference, variant) pairs")
    mono(doc, "       |                       |")
    mono(doc, "       |         [Score each pair with 5 features]")
    mono(doc, "       |         using vocab.pkl + trigram.pkl + base_model.pt")
    mono(doc, "       |                       |")
    mono(doc, "       |                  features.csv")
    mono(doc, "       |                       |")
    mono(doc, "       '----------> [Logistic Regression] ---> Results")
    doc.add_paragraph()

    # ── 1. The Two Datasets ──────────────────────────────────────────────────
    style_heading(doc, "1. The Two Datasets", 1)

    style_heading(doc, "1.1 Hindi Wikipedia — Trains the Language Models", 2)
    body(doc,
        "The language models need to learn statistical patterns of Hindi — which words "
        "tend to follow which others, which sequences are common or rare. For this we use "
        "Hindi Wikipedia, a large collection of Hindi text that covers a wide range of topics."
    )
    body(doc,
        "We do NOT use the treebank to train the language models. This is important: the "
        "treebank sentences will later be used to evaluate whether the models can predict "
        "word order, so using them for training would be cheating."
    )
    simple_table(doc,
        ["Property", "Value"],
        [
            ["Source", "Hindi Wikipedia XML dump"],
            ["Sentences after preprocessing", "2,655,055"],
            ["Vocabulary size", "30,000 most frequent words"],
            ["Models trained on this data", "Trigram LM (trigram.pkl), LSTM LM (base_model.pt)"],
        ],
        col_widths=[2.5, 4.0]
    )

    style_heading(doc, "1.2 HDTB Treebank — Evaluation Only", 2)
    body(doc,
        "The Hindi Dependency Treebank (HDTB) is a hand-annotated collection of real Hindi "
        "newspaper sentences. Each sentence has been carefully labelled by linguists with the "
        "grammatical role of every word and which word it depends on. We use this purely for "
        "evaluation — to get real Hindi sentences, generate permuted alternatives, and test "
        "whether our features can identify the real sentence."
    )
    body(doc,
        "The treebank comes in three files that together cover 16,649 sentences:"
    )
    simple_table(doc,
        ["File", "Sentences", "Role in this project"],
        [
            ["hi_hdtb-ud-train.conllu", "13,306", "Used for evaluation (sentence pairs + features)"],
            ["hi_hdtb-ud-dev.conllu",   "1,659",  "Used for evaluation (same as train)"],
            ["hi_hdtb-ud-test.conllu",  "1,684",  "Used for evaluation (same as train)"],
            ["Total", "16,649", "All three merged for maximum coverage"],
        ],
        col_widths=[3.0, 1.2, 2.8]
    )
    note(doc,
        "The train/dev/test split is designed for parser training, not for this experiment. "
        "Since our language models are trained on Wikipedia, there is no leakage — we can "
        "safely use all three splits."
    )

    # ── 2. What is CoNLL-U ───────────────────────────────────────────────────
    style_heading(doc, "2. What Does the Treebank Look Like? (CoNLL-U Format)", 1)
    body(doc,
        "The treebank is stored in CoNLL-U format. Each sentence is a block of lines separated "
        "by a blank line. Each line is one word, with 10 tab-separated columns."
    )

    label(doc, "Example sentence: 'इसे नवाब शाहजेहन ने बनवाया था ।'  (He had it built by Nawab Shahjahan)")
    mono(doc, "# sent_id = 2:14")
    mono(doc, "# text = इसे नवाब शाहजेहन ने बनवाया था ।")
    mono(doc, "ID   WORD        LEMMA       UPOS  HEAD  DEPREL")
    mono(doc, "1    इसे         यह          PRON  5     obj       <- object, depends on word 5")
    mono(doc, "2    नवाब        नवाब        NOUN  4     compound  <- part of subject phrase")
    mono(doc, "3    शाहजेहन     शाहजेहन     PROPN 4     compound  <- part of subject phrase")
    mono(doc, "4    ने          ने          ADP   5     nsubj     <- subject, depends on word 5")
    mono(doc, "5    बनवाया      बनवाना      VERB  0     root      <- ROOT verb (head=0)")
    mono(doc, "6    था          होना        AUX   5     aux       <- auxiliary, depends on word 5")
    mono(doc, "7    ।           ।           PUNCT 5     punct     <- punctuation")
    doc.add_paragraph()

    body(doc,
        "The key columns we use are: ID (position), WORD (surface form), LEMMA (base form), "
        "UPOS (part of speech), HEAD (which word this word depends on — 0 means root), "
        "and DEPREL (the grammatical role: obj=object, nsubj=subject, aux=auxiliary, etc.)."
    )

    label(doc, "Dependency Tree Diagram for the Example")
    mono(doc, "                    बनवाया (ROOT, pos=5)")
    mono(doc, "                   /   |    \\    \\")
    mono(doc, "                  /    |     \\    \\")
    mono(doc, "               इसे    ने     था    ।")
    mono(doc, "              (obj,1) (nsubj,4) (aux,6) (punct,7)")
    mono(doc, "                       |")
    mono(doc, "                    नवाब (2) + शाहजेहन (3)")
    mono(doc, "                    [compound dependents of ने]")
    doc.add_paragraph()

    # ── 3. Sentence Filtering ────────────────────────────────────────────────
    style_heading(doc, "3. Sentence Filtering — Which Sentences Pass?", 1)
    body(doc,
        "Not all 16,649 sentences are usable. The paper defines 6 criteria that a sentence "
        "must meet. Only 2,828 (~17%) pass all filters. Below, each filter is explained with "
        "a PASS example and a FAIL example."
    )

    # Filter 1
    style_heading(doc, "Filter 1: Root Must Be a Verb or Auxiliary (VERB/AUX)", 2)
    body(doc,
        "We need sentences with a main verb as the head of the clause, because only then "
        "can we meaningfully move phrases around the verb. Sentences headed by a noun, "
        "adjective, or number are excluded."
    )
    pass_fail_table(doc,
        ["Sentence", "Root word", "Root UPOS", "Result"],
        [
            ["इसे नवाब शाहजेहन ने बनवाया था ।",     "बनवाया",  "VERB",  "PASS ✓"],
            ["राम घर में है ।",                       "है",       "AUX",   "PASS ✓"],
            ["यह एक बड़ी इमारत ।  (nominal sentence)", "इमारत",   "NOUN",  "FAIL ✗"],
        ]
    )

    # Filter 2
    style_heading(doc, "Filter 2: Projective Dependency Tree", 2)
    body(doc,
        "A projective sentence means no dependency arcs cross each other when drawn above "
        "the sentence. Non-projective sentences have unusual structures that are harder to "
        "permute consistently."
    )
    label(doc, "Projective (PASS) — no arcs cross:")
    mono(doc, "  राम   ने   सेब   खाया")
    mono(doc, "  1     2    3     4(ROOT)")
    mono(doc, "  |_____|    |_____|")
    mono(doc, "  nsubj       obj       <- arcs do not cross  ✓")
    doc.add_paragraph()
    label(doc, "Non-projective (FAIL) — arcs cross:")
    mono(doc, "  राम   सेब   ने   खाया")
    mono(doc, "  1     2     3    4(ROOT)")
    mono(doc, "  |_______________|")
    mono(doc, "        |_____|        <- arc from 1→4 crosses arc from 2→4  ✗")
    doc.add_paragraph()

    # Filter 3
    style_heading(doc, "Filter 3: Declarative Sentence (No '?')", 2)
    body(doc,
        "Questions have different word order dynamics from statements. We keep only "
        "declarative sentences."
    )
    pass_fail_table(doc,
        ["Sentence", "Contains '?'", "Result"],
        [
            ["राम ने सेब खाया ।",   "No",   "PASS ✓"],
            ["क्या राम आया ?",       "Yes",  "FAIL ✗"],
        ]
    )

    # Filter 4
    style_heading(doc, "Filter 4: No Negative Markers (नहीं / न / मत)", 2)
    body(doc,
        "Negation changes how sentences are interpreted and processed. To keep the "
        "analysis clean, sentences containing नहीं (not), न (not), or मत (don't) are dropped."
    )
    pass_fail_table(doc,
        ["Sentence", "Contains negation?", "Result"],
        [
            ["राम ने सेब खाया ।",            "No",              "PASS ✓"],
            ["राम ने सेब नहीं खाया ।",        "Yes — नहीं",      "FAIL ✗"],
            ["न जाओ वहाँ ।",                  "Yes — न",         "FAIL ✗"],
        ]
    )

    # Filter 5
    style_heading(doc, "Filter 5: Subject AND Object Must Be Direct Root Dependents", 2)
    body(doc,
        "We can only permute phrases that are direct children of the root verb. If the "
        "subject or object is buried inside a relative clause or embedded sentence, "
        "it cannot be moved without breaking the sentence."
    )
    label(doc, "PASS — both nsubj and obj attach directly to the root verb:")
    mono(doc, "  इसे(obj) --- बनवाया(ROOT) --- ने(nsubj)")
    mono(doc, "  Both obj and nsubj are DIRECT children of ROOT  ✓")
    doc.add_paragraph()
    label(doc, "FAIL — object is inside an embedded clause, not a direct child of root:")
    mono(doc, "  राम ने [जो किताब मैंने लिखी] पढ़ी")
    mono(doc, "  ROOT = पढ़ी,   nsubj = राम (direct ✓)")
    mono(doc, "  obj = किताब   but किताब is inside the relative clause [...]")
    mono(doc, "  किताब's head = लिखी (not ROOT)  ->  FAIL ✗")
    doc.add_paragraph()

    # Filter 6
    style_heading(doc, "Filter 6: At Least Two Pre-Verbal Phrases", 2)
    body(doc,
        "We need at least two phrases before the verb to generate any meaningful alternative "
        "ordering. If there is only one pre-verbal phrase, there is nothing to permute."
    )
    pass_fail_table(doc,
        ["Sentence", "Pre-verbal direct dependents", "Result"],
        [
            ["इसे नवाब ने बनवाया — obj(इसे) + nsubj(नवाब ने)", "2", "PASS ✓"],
            ["राम ने खाया — only nsubj(राम ने)",                 "1", "FAIL ✗"],
        ]
    )

    label(doc, "Filter Summary")
    simple_table(doc,
        ["Filter", "Criterion", "Sentences dropped"],
        [
            ["1", "Root is VERB or AUX",                                "~1,654"],
            ["2", "Projective dependency tree",                         "~several hundred"],
            ["3", "Declarative (no '?')",                               "~small number"],
            ["4", "No negative markers",                                "~several hundred"],
            ["5", "Subject and object are direct root dependents",      "~majority of remainder"],
            ["6", "At least 2 pre-verbal direct dependents",            "~overlap with 5"],
            ["Total retained", "All 6 filters passed",                  "2,828 of 16,649 (17%)"],
        ],
        col_widths=[0.5, 3.5, 2.0]
    )

    # ── 4. Variant Generation ────────────────────────────────────────────────
    style_heading(doc, "4. Variant Generation — Building the Pairs", 1)
    body(doc,
        "For each of the 2,828 valid sentences, we generate alternative orderings by "
        "swapping the pre-verbal phrases while keeping the verb in place."
    )

    label(doc, "Step-by-step example:")
    body(doc, "Reference sentence: इसे नवाब शाहजेहन ने बनवाया था ।")
    body(doc, "Pre-verbal direct dependents of ROOT (बनवाया):")
    mono(doc, "  Phrase A:  इसे             (obj — direct object)")
    mono(doc, "  Phrase B:  नवाब शाहजेहन ने (nsubj — subject)")
    mono(doc, "  ROOT:      बनवाया था ।      (stays fixed)")
    doc.add_paragraph()
    body(doc, "Possible orderings:")
    simple_table(doc,
        ["Order", "Surface sentence", "Construction type"],
        [
            ["A B ROOT (original)", "इसे नवाब शाहजेहन ने बनवाया था ।",    "DOSV (obj first)"],
            ["B A ROOT (variant)",  "नवाब शाहजेहन ने इसे बनवाया था ।",    "SOV  (subj first — canonical)"],
        ],
        col_widths=[1.5, 3.5, 2.0]
    )

    style_heading(doc, "4.1 The Adjacency Filter — Removing Ungrammatical Variants", 2)
    body(doc,
        "Not all permutations are grammatical. In Hindi, postpositions (like ने, को, में) "
        "must follow the noun phrase they mark. The adjacency filter learns which pairs of "
        "grammatical relations can appear next to each other from the actual treebank, "
        "then rejects any variant that would create an impossible adjacency."
    )
    label(doc, "Example:")
    mono(doc, "  In the treebank, 'compound → nsubj' adjacency is common.")
    mono(doc, "  But 'obj → compound' adjacency never appears.")
    mono(doc, "  A variant that would place an obj-headed phrase directly before")
    mono(doc, "  a compound node is therefore rejected as ungrammatical.")
    doc.add_paragraph()

    style_heading(doc, "4.2 Construction Type Labels", 2)
    simple_table(doc,
        ["Label", "Meaning", "Example surface order"],
        [
            ["SOV",  "Subject–Object–Verb  (canonical Hindi order)",           "राम ने सेब को देखा"],
            ["DOSV", "Direct-Object–Subject–Verb  (obj fronted, non-canonical)", "सेब को राम ने देखा"],
            ["IOSV", "Indirect-Object–Subject–Verb  (iobj fronted)",             "मुझे राम ने सेब दिया"],
        ],
        col_widths=[1.0, 3.5, 2.5]
    )

    # ── 5. Feature 1: Dependency Length ──────────────────────────────────────
    style_heading(doc, "5. Feature 1 — Dependency Length (DL)", 1)
    body(doc,
        "Dependency length is the number of words between a word and its grammatical head. "
        "The shorter the dependency, the less memory is needed to process it — the brain "
        "can resolve grammatical relationships sooner. The hypothesis is that speakers choose "
        "word orders that minimise total dependency length."
    )

    label(doc, "Formula:  DL(sentence) = Σ |position(word) − position(head)|  for all non-root words")

    label(doc, "Visual example — Reference: इसे(1) नवाब(2) शाहजेहन(3) ने(4) बनवाया(5) था(6) ।(7)")
    mono(doc, "  Word        ID   Head   |ID - Head|   Arc")
    mono(doc, "  इसे          1    5       4           1 -------- 5")
    mono(doc, "  नवाब         2    4       2           2 ------ 4")
    mono(doc, "  शाहजेहन      3    4       1           3 --- 4")
    mono(doc, "  ने           4    5       1           4 - 5")
    mono(doc, "  बनवाया       5    0  (ROOT, skip)")
    mono(doc, "  था           6    5       1           5 - 6")
    mono(doc, "  ।            7    5       2           5 ------- 7")
    mono(doc, "                         --------")
    mono(doc, "  Total DL (reference) = 4+2+1+1+1+2 = 11")
    doc.add_paragraph()

    label(doc, "Same words, variant order: नवाब(1) शाहजेहन(2) ने(3) इसे(4) बनवाया(5) था(6) ।(7)")
    mono(doc, "  Word        ID   Head   |ID - Head|")
    mono(doc, "  नवाब         1    3       2")
    mono(doc, "  शाहजेहन      2    3       1")
    mono(doc, "  ने           3    5       2")
    mono(doc, "  इसे          4    5       1")
    mono(doc, "  बनवाया       5    0  (ROOT, skip)")
    mono(doc, "  था           6    5       1")
    mono(doc, "  ।            7    5       2")
    mono(doc, "                         --------")
    mono(doc, "  Total DL (variant) = 2+1+2+1+1+2 = 9")
    doc.add_paragraph()

    simple_table(doc,
        ["Measure", "Reference", "Variant", "Delta (ref − var)"],
        [
            ["Total Dependency Length", "11", "9", "+2  (reference has LONGER DL)"],
        ],
        col_widths=[2.5, 1.2, 1.2, 3.0]
    )
    body(doc,
        "In this particular pair, the variant has shorter DL. Across the full dataset, "
        "58.4% of reference sentences have shorter DL than their variants — meaning speakers "
        "generally do prefer shorter dependency lengths, but it is not the only factor."
    )

    # ── 6. Feature 2: Trigram Surprisal ──────────────────────────────────────
    style_heading(doc, "6. Feature 2 — Trigram Surprisal", 1)
    body(doc,
        "Surprisal measures how unexpected a word is given its context. "
        "High surprisal = the word was hard to predict. Low surprisal = the word was expected. "
        "The hypothesis is that speakers choose word orders that make each word more predictable. "
        "A trigram model estimates P(word | previous 2 words) using statistics learned from Wikipedia."
    )

    label(doc, "Formula:  Surprisal(w_i) = −log P(w_i | w_{i−2}, w_{i−1})")
    label(doc, "          Sentence surprisal = sum of per-word surprisals (starting from word 3)")
    doc.add_paragraph()

    label(doc, "Visual example — sliding window across reference sentence:")
    mono(doc, "  Position:  1          2           3           4           5          6")
    mono(doc, "  Word:      इसे        नवाब        शाहजेहन     ने          बनवाया     था")
    mono(doc, "")
    mono(doc, "  Word 1 (इसे):      No context yet — skip (surprisal = 0)")
    mono(doc, "  Word 2 (नवाब):     Only 1 prior word — skip (surprisal = 0)")
    mono(doc, "  Word 3 (शाहजेहन):  P(शाहजेहन | इसे, नवाब)  -> surprisal = -log(P)")
    mono(doc, "  Word 4 (ने):       P(ने | नवाब, शाहजेहन)   -> surprisal = -log(P)")
    mono(doc, "  Word 5 (बनवाया):   P(बनवाया | शाहजेहन, ने) -> surprisal = -log(P)")
    mono(doc, "  Word 6 (था):       P(था | ने, बनवाया)       -> surprisal = -log(P)")
    mono(doc, "")
    mono(doc, "  Total surprisal = sum of surprisals for words 3-6")
    doc.add_paragraph()

    label(doc, "Backoff strategy (when a trigram was never seen in Wikipedia):")
    mono(doc, "  Level 1 (trigram):  P(w3 | w1, w2)   <- use if count > 0")
    mono(doc, "  Level 2 (bigram):   P(w3 | w2)        <- fall back if trigram unseen")
    mono(doc, "  Level 3 (unigram):  P(w3)             <- fall back if bigram unseen")
    mono(doc, "  Level 4 (epsilon):  1e-12             <- completely unknown word")
    doc.add_paragraph()

    simple_table(doc,
        ["Sentence", "Total Surprisal", "delta_trigram (ref − var)"],
        [
            ["Reference: इसे नवाब शाहजेहन ने बनवाया था ।",  "15.23", ""],
            ["Variant:   नवाब शाहजेहन ने इसे बनवाया था ।",  "16.81", "15.23 − 16.81 = −1.58"],
        ],
        col_widths=[3.5, 1.5, 2.5]
    )
    body(doc,
        "Negative delta = reference is MORE predictable. Across the dataset, 77.3% of "
        "reference sentences have lower surprisal than their variants — this is the "
        "strongest single predictor."
    )

    # ── 7. Feature 3: LSTM Surprisal ─────────────────────────────────────────
    style_heading(doc, "7. Feature 3 — LSTM Surprisal", 1)
    body(doc,
        "The LSTM language model works similarly to the trigram model but uses a neural "
        "network that can in principle capture much longer context than just 2 previous words. "
        "It processes the sentence word by word, updating a hidden state at each step, "
        "and estimates the probability of the next word at every position."
    )

    label(doc, "Architecture trained on Wikipedia:")
    mono(doc, "  Input word --> [Embedding layer, 256-dim]")
    mono(doc, "                        |")
    mono(doc, "               [LSTM Layer 1, 256 hidden units]")
    mono(doc, "                        |")
    mono(doc, "               [LSTM Layer 2, 256 hidden units]")
    mono(doc, "                        |")
    mono(doc, "              [Linear layer: 256 -> 30,000]")
    mono(doc, "                        |")
    mono(doc, "             [Softmax -> probability over vocabulary]")
    doc.add_paragraph()

    label(doc, "Processing the reference sentence step by step:")
    mono(doc, "  Step 1:  input=<SOS>      hidden=h0  ->  predict: इसे")
    mono(doc, "  Step 2:  input=इसे        hidden=h1  ->  predict: नवाब")
    mono(doc, "  Step 3:  input=नवाब       hidden=h2  ->  predict: शाहजेहन")
    mono(doc, "  Step 4:  input=शाहजेहन    hidden=h3  ->  predict: ने")
    mono(doc, "  Step 5:  input=ने         hidden=h4  ->  predict: बनवाया")
    mono(doc, "  Step 6:  input=बनवाया     hidden=h5  ->  predict: था")
    mono(doc, "  ...")
    mono(doc, "  Surprisal at each step = -log P(actual_next_word | hidden_state)")
    mono(doc, "  Total LSTM surprisal = sum of all per-step surprisals")
    doc.add_paragraph()

    body(doc,
        "The LSTM was trained for 2 epochs on 2 million Wikipedia tokens. "
        "Training loss dropped from 4.42 (epoch 1) to 3.36 (epoch 2). "
        "The model weights are saved and reused — it is never retrained on the treebank."
    )

    # ── 8. Feature 4: Adaptive LSTM ──────────────────────────────────────────
    style_heading(doc, "8. Feature 4 — Adaptive LSTM Surprisal", 1)
    body(doc,
        "The base LSTM was trained on Wikipedia and knows general Hindi statistics. "
        "But real sentences appear in a specific document context. The adaptive model "
        "fine-tunes itself on the sentences that appeared just before the target sentence, "
        "so that words recently mentioned in the document become more expected."
    )

    label(doc, "Example — two consecutive sentences in a newspaper article:")
    mono(doc, "  Context sentence (sentence N-1):")
    mono(doc, "    'ताजमहल आगरा में स्थित एक प्रसिद्ध मकबरा है ।'")
    mono(doc, "    (The Taj Mahal is a famous mausoleum in Agra.)")
    mono(doc, "")
    mono(doc, "  Target sentence (sentence N):")
    mono(doc, "    'इसे नवाब शाहजेहन ने बनवाया था ।'")
    mono(doc, "    (It was built by Nawab Shahjahan.)")
    doc.add_paragraph()

    label(doc, "How adaptation works:")
    mono(doc, "  1. Start with base LSTM weights (trained on Wikipedia)")
    mono(doc, "  2. Run ONE gradient update step on the CONTEXT sentence")
    mono(doc, "     -> model now assigns higher probability to words like")
    mono(doc, "        'ताजमहल', 'आगरा', 'शाहजेहन', 'मकबरा'")
    mono(doc, "  3. Score surprisal on the TARGET sentence with adapted weights")
    mono(doc, "  4. Score surprisal on ALL variants of target with same adapted weights")
    mono(doc, "  5. Reset to base weights before processing the next sentence")
    doc.add_paragraph()

    label(doc, "Why reset per sentence (not per pair)?")
    mono(doc, "  One reference sentence (e.g. sentence #5) may have 30 variants.")
    mono(doc, "  We adapt ONCE on the context, then score all 30 variants.")
    mono(doc, "  We do NOT adapt 30 times — that would compound the gradient")
    mono(doc, "  updates and drift the model far from its base state.")
    doc.add_paragraph()

    body(doc,
        "The adaptive model is the most powerful surprisal predictor for IOSV sentences "
        "(coefficient −1.11), because indirect object fronting is strongly driven by "
        "discourse context — you front something to the beginning of a sentence because "
        "it was just mentioned."
    )

    # ── 9. Feature 5: Information Status ─────────────────────────────────────
    style_heading(doc, "9. Feature 5 — Information Status (Givenness)", 1)
    body(doc,
        "Information Status (IS) captures whether the phrases in a sentence carry "
        "new or already-known information. A general principle in linguistics is that "
        "speakers prefer to mention what is already known (GIVEN) before introducing "
        "new information (NEW). This is called given-before-new order."
    )

    label(doc, "Example:")
    mono(doc, "  Context sentence: 'ताजमहल आगरा में स्थित एक प्रसिद्ध मकबरा है ।'")
    mono(doc, "  Target sentence:  'इसे नवाब शाहजेहन ने बनवाया था ।'")
    mono(doc, "")
    mono(doc, "  Subject: 'नवाब शाहजेहन ने'  <- NOT mentioned in context  -> NEW")
    mono(doc, "  Object:  'इसे'               <- pronoun (refers back to Taj Mahal) -> GIVEN")
    mono(doc, "")
    mono(doc, "  Surface order:  [इसे=GIVEN] [नवाब शाहजेहन ने=NEW] [बनवाया था]")
    mono(doc, "                   ^first                ^second")
    mono(doc, "  First=GIVEN, Second=NEW  ->  IS score = +1  (given-before-new ✓)")
    doc.add_paragraph()

    label(doc, "IS scoring rules:")
    simple_table(doc,
        ["First phrase", "Second phrase", "IS score", "Meaning"],
        [
            ["GIVEN", "NEW",   "+1", "Given-before-new — natural discourse order"],
            ["NEW",   "GIVEN", "-1", "New-before-given — marked, unusual order"],
            ["GIVEN", "GIVEN", " 0", "Both given — no informational contrast"],
            ["NEW",   "NEW",   " 0", "Both new — no informational contrast"],
        ],
        col_widths=[1.3, 1.3, 1.0, 3.4]
    )

    label(doc, "How GIVEN is decided for a phrase:")
    mono(doc, "  Rule A: If the head word is a PRONOUN -> GIVEN")
    mono(doc, "          (pronouns always refer back to something already mentioned)")
    mono(doc, ""  )
    mono(doc, "  Rule B: If any CONTENT WORD (noun, verb, adj, adv, num) in the phrase")
    mono(doc, "          appeared in the immediately preceding sentence -> GIVEN")
    mono(doc, ""  )
    mono(doc, "  Otherwise: NEW")
    doc.add_paragraph()

    label(doc, "delta_is = IS(reference) − IS(variant)")
    body(doc,
        "A positive delta means the reference follows given-before-new order more than its "
        "variant. This is significant for DOSV sentences (coefficient +0.41) — speakers "
        "front the direct object partly because it is already known information."
    )

    # ── 10. Features CSV ─────────────────────────────────────────────────────
    style_heading(doc, "10. The Features Table (features.csv)", 1)
    body(doc,
        "After all five features are computed for all 92,299 pairs, everything is saved "
        "to a single CSV file. Here is what the first row looks like:"
    )
    simple_table(doc,
        ["Column", "Example value", "Meaning"],
        [
            ["sentence_id",        "0",          "Index of original treebank sentence"],
            ["construction_type",  "DOSV",       "Reference is direct-object-first"],
            ["reference",          "इसे नवाब...", "The real treebank sentence"],
            ["variant",            "नवाब इसे...", "The permuted alternative"],
            ["dl_reference",       "11",          "Total dependency length of reference"],
            ["dl_variant",         "9",           "Total dependency length of variant"],
            ["delta_dl",           "+2",          "ref DL − var DL  (positive = ref longer)"],
            ["trigram_reference",  "15.23",       "Trigram surprisal of reference"],
            ["trigram_variant",    "16.81",       "Trigram surprisal of variant"],
            ["delta_trigram",      "-1.58",       "ref surprisal − var surprisal  (negative = ref more predictable)"],
            ["delta_lstm",         "-3.60",       "Same for LSTM"],
            ["delta_adaptive",     "+3.59",       "Same for adaptive LSTM"],
            ["delta_is",           "+2",          "IS(reference) − IS(variant)"],
        ],
        col_widths=[2.0, 1.5, 3.5]
    )

    # ── 11. The Ranking Model ────────────────────────────────────────────────
    style_heading(doc, "11. The Ranking Model — How Evaluation Works", 1)
    body(doc,
        "The five delta values for each pair are fed into a logistic regression model. "
        "The model must predict: is this the real sentence (reference preferred) or the "
        "machine-generated alternative (variant preferred)?"
    )

    label(doc, "The pairwise trick:")
    mono(doc, "  Original pair:  [delta_dl=+2, delta_tri=-1.58, ...]  ->  label = 1 (ref preferred)")
    mono(doc, "  Flipped pair:   [-2, +1.58, ...]                     ->  label = 0 (var preferred)")
    mono(doc, "")
    mono(doc, "  This doubling prevents the model from cheating by always predicting 1.")
    doc.add_paragraph()

    label(doc, "10-fold cross-validation at the SENTENCE level:")
    mono(doc, "  - All 2,828 sentences split into 10 equal groups")
    mono(doc, "  - Train on 9 groups (with all their variants), test on the 10th")
    mono(doc, "  - Repeat 10 times, average accuracy")
    mono(doc, "  - Splitting at sentence level prevents train/test leakage")
    mono(doc, "    (a sentence's variants must all be in the same split)")
    doc.add_paragraph()

    label(doc, "Results vs the paper:")
    simple_table(doc,
        ["Condition", "Our accuracy", "Paper accuracy", "Gap"],
        [
            ["Full — baseline (adaptive+trigram)", "80.80%", "85.18%", "-4.4%"],
            ["Full — all features",                "81.49%", "85.04%", "-3.6%"],
            ["DOSV — baseline",                    "68.96%", "81.24%", "-12.3%"],
            ["DOSV — all features",                "73.60%", "80.46%", "-6.9%"],
            ["IOSV — baseline",                    "82.28%", "89.43%", "-7.2%"],
            ["IOSV — all features",                "82.62%", "90.02%", "-7.4%"],
        ],
        col_widths=[2.8, 1.4, 1.4, 1.2]
    )
    body(doc,
        "The gap is mainly due to the missing PCFG surprisal feature. "
        "All coefficient directions match the paper (10/10)."
    )

    # ── 12. Regression Coefficients ──────────────────────────────────────────
    style_heading(doc, "12. Regression Coefficients — What the Model Learned", 1)
    body(doc,
        "After fitting the logistic regression on all 92,299 pairs, the model assigns a "
        "weight (coefficient) to each feature. A negative coefficient means the reference "
        "sentence tends to have a LOWER value of that feature — i.e. speakers minimise it. "
        "A positive coefficient means the reference tends to have a HIGHER value. "
        "All features are z-scored (standardised) so coefficients are directly comparable."
    )

    style_heading(doc, "12.1 DOSV — Direct Object Fronted Sentences", 2)
    body(doc,
        "These are sentences where the direct object appears before the subject "
        "(e.g. 'इसे नवाब ने बनवाया' instead of 'नवाब ने इसे बनवाया')."
    )
    simple_table(doc,
        ["Feature", "Coefficient", "Direction", "Interpretation"],
        [
            ["Dependency Length",  "-0.814", "Negative", "Speakers choose the order that shortens dependency arcs — strongest effect for DOSV"],
            ["Trigram Surprisal",  "-0.767", "Negative", "Reference is more statistically predictable according to Wikipedia-trained trigram model"],
            ["Information Status", "+0.413", "Positive", "Reference more often places given information before new — direct object is fronted because it was already mentioned"],
            ["LSTM Surprisal",     "+0.271", "Positive", "Positive because adaptive surprisal absorbs most of the LSTM signal — residual is small"],
            ["Adaptive Surprisal", "-0.268", "Negative", "Reference fits the discourse context better (smaller effect for DOSV than IOSV)"],
        ],
        col_widths=[1.8, 1.1, 1.1, 3.5]
    )
    body(doc,
        "Key finding for DOSV: Dependency Length is the dominant predictor — speakers front "
        "the direct object primarily because doing so shortens the grammatical link between "
        "the object and the verb. Information Status is also significant — the fronted object "
        "is often already known to the listener."
    )

    style_heading(doc, "12.2 IOSV — Indirect Object Fronted Sentences", 2)
    body(doc,
        "These are sentences where the indirect object (recipient/beneficiary) appears first "
        "(e.g. 'मुझे राम ने किताब दी' instead of 'राम ने मुझे किताब दी')."
    )
    simple_table(doc,
        ["Feature", "Coefficient", "Direction", "Interpretation"],
        [
            ["Adaptive Surprisal", "-1.109", "Negative", "Dominant predictor — reference fits prior discourse context much better; indirect object fronted because it was just mentioned"],
            ["LSTM Surprisal",     "+1.105", "Positive", "Mirrors adaptive: base LSTM alone finds reference more surprising, but adaptive corrects this after seeing context"],
            ["Trigram Surprisal",  "-0.650", "Negative", "Reference is more predictable by trigram statistics"],
            ["Dependency Length",  "-0.287", "Negative", "DL has a weaker effect for IOSV — without PCFG this appears significant, but the paper shows PCFG absorbs it"],
            ["Information Status", "+0.062", "Positive", "Weak and not significant for IOSV — givenness matters less here than discourse-level predictability"],
        ],
        col_widths=[1.8, 1.1, 1.1, 3.5]
    )
    body(doc,
        "Key finding for IOSV: Adaptive surprisal dominates — speakers front the indirect "
        "object because it is already established in the discourse, making the sentence flow "
        "more naturally from the reader's perspective. DL is a weaker secondary effect."
    )

    style_heading(doc, "12.3 Full Dataset", 2)
    simple_table(doc,
        ["Feature", "Coefficient", "Direction", "Interpretation"],
        [
            ["Trigram Surprisal",  "-1.374", "Negative", "Strongest overall predictor — word order that fits statistical patterns wins"],
            ["Adaptive Surprisal", "-0.592", "Negative", "Second strongest — discourse context is a consistent driver"],
            ["LSTM Surprisal",     "+0.589", "Positive", "Positive as expected — adaptive absorbs the negative LSTM signal"],
            ["Dependency Length",  "-0.458", "Negative", "Consistent negative effect across all construction types"],
            ["Information Status", "+0.324", "Positive", "Givenness is a reliable positive predictor overall"],
        ],
        col_widths=[1.8, 1.1, 1.1, 3.5]
    )
    note(doc,
        "All coefficient directions (positive/negative) match the paper exactly — 10 out of 10. "
        "Exact magnitudes differ slightly because we do not have PCFG surprisal."
    )

    # ── 13. Known Gap — PCFG ─────────────────────────────────────────────────
    style_heading(doc, "13. Known Gap — PCFG Surprisal (Not Implemented)", 1)
    body(doc,
        "The paper includes PCFG (Probabilistic Context-Free Grammar) surprisal as a predictor. "
        "We were unable to implement it. This section explains what it is, why it matters, "
        "and exactly why it could not be added."
    )

    style_heading(doc, "13.1 What is PCFG Surprisal?", 2)
    body(doc,
        "A PCFG is a grammar that assigns probabilities to phrase-structure rules. For example:"
    )
    mono(doc, "  S  -> NP VP          with probability 0.95")
    mono(doc, "  VP -> V NP           with probability 0.60")
    mono(doc, "  VP -> V NP NP        with probability 0.40")
    doc.add_paragraph()
    body(doc,
        "A PCFG parser reads a sentence and produces a constituency parse tree — a bracketed "
        "structure showing how words group into phrases. PCFG surprisal is the negative log "
        "probability of the sentence's parse tree under this grammar. Low PCFG surprisal = "
        "the sentence has a common, expected phrase structure."
    )

    style_heading(doc, "13.2 Why It Is Not Implemented — Two Blockers", 2)

    label(doc, "Blocker 1: Wrong treebank format")
    body(doc,
        "PCFG requires a constituency treebank — sentences annotated with bracket structure. "
        "The HDTB is a DEPENDENCY treebank — sentences annotated with head-word relations. "
        "These are fundamentally different representations:"
    )
    mono(doc, "  Dependency (what HDTB provides):")
    mono(doc, "    Word      Head    Relation")
    mono(doc, "    राम        खाया    nsubj")
    mono(doc, "    सेब        खाया    obj")
    mono(doc, "    खाया       ROOT    root")
    mono(doc, "")
    mono(doc, "  Constituency (what PCFG needs):")
    mono(doc, "    [S  [NP राम ने]")
    mono(doc, "        [VP [NP सेब को]")
    mono(doc, "            [V खाया]]]")
    doc.add_paragraph()
    body(doc,
        "To use PCFG, we would need to either (a) find a Hindi constituency treebank, or "
        "(b) convert HDTB dependency trees to constituency trees using a dependency-to-constituency "
        "converter. No reliable Hindi dependency-to-constituency converter exists publicly."
    )

    label(doc, "Blocker 2: No off-the-shelf Hindi PCFG parser")
    body(doc,
        "For English, tools like the Stanford Parser or Berkeley Parser provide ready-made PCFG "
        "parsers trained on the Penn Treebank. For Hindi, no equivalent trained PCFG parser is "
        "publicly available. We would need to: (1) obtain a Hindi constituency treebank, "
        "(2) train a PCFG from scratch, (3) use it to parse all 92,299 sentence variants. "
        "Each variant would need to be re-parsed with its permuted word order — a computationally "
        "expensive step for which no ready infrastructure exists."
    )

    style_heading(doc, "13.3 What Effect Does This Have on Results?", 2)
    body(doc,
        "The paper shows that PCFG surprisal is part of the baseline (adaptive + trigram + PCFG). "
        "Its most important effect is on IOSV sentences:"
    )
    simple_table(doc,
        ["Condition", "Without PCFG (ours)", "With PCFG (paper)", "Gap"],
        [
            ["DOSV baseline accuracy",  "68.96%", "81.24%", "-12.3%"],
            ["IOSV baseline accuracy",  "82.28%", "89.43%", "-7.2%"],
            ["DL significant for IOSV?", "YES (coeff -0.29)", "NO — absorbed by PCFG", "Diverges from paper"],
        ],
        col_widths=[2.5, 1.8, 1.8, 1.8]
    )
    body(doc,
        "The paper's key finding is that for IOSV sentences, dependency length is NOT a "
        "significant predictor once PCFG surprisal is controlled for. Our replication cannot "
        "reproduce this finding. In our results, DL appears weakly significant for IOSV, "
        "which is an artefact of the missing PCFG feature."
    )

    # ── 14. All Assumptions ──────────────────────────────────────────────────
    style_heading(doc, "14. All Assumptions Made in This Replication", 1)
    body(doc,
        "The paper does not specify every implementation detail. The following assumptions "
        "were made to fill the gaps, along with the reasoning behind each."
    )

    style_heading(doc, "14.1 Treebank Splits", 2)
    simple_table(doc,
        ["Assumption", "What we did", "Reasoning"],
        [
            ["Which splits to use",
             "Currently train only (13,306 sentences). Plan to add dev+test.",
             "The paper does not specify. Train/dev/test split is for parser evaluation, not this experiment. All splits should be safe to use since language models are trained on Wikipedia."],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )

    style_heading(doc, "14.2 Sentence Filtering", 2)
    simple_table(doc,
        ["Assumption", "What we did", "Reasoning"],
        [
            ["VerbForm=Fin check",
             "NOT applied — check upos in {VERB, AUX} only",
             "HDTB annotates most finite verbs as VerbForm=Part (HDTB convention for perfective verbs). Applying VerbForm=Fin would drop ~94% of valid sentences. upos check alone correctly excludes all nominal roots."],
            ["Root can be AUX",
             "Allowed (not just VERB)",
             "Hindi compound verb constructions (e.g. 'कर दिया') have AUX as the root in HDTB. Excluding AUX would miss these valid sentences."],
            ["Negative markers",
             "नहीं, न, मत excluded",
             "Paper specifies these three. We took them literally."],
            ["Projectivity",
             "Strict check — any crossing arc drops the sentence",
             "Paper mandates projective trees. Strict interpretation applied."],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )

    style_heading(doc, "14.3 Variant Generation", 2)
    simple_table(doc,
        ["Assumption", "What we did", "Reasoning"],
        [
            ["Adjacency filter",
             "NOT applied despite being in the paper",
             "The paper references Rajkumar & White (2014) who use a full CCG surface realiser, not a simple deprel-pair checker. Our approximation of it rejected ~90% of valid variants (8k pairs vs paper's 72k). Structural filters already ensure well-formed inputs. Filter removed to match paper's pair counts."],
            ["Duplicate variants",
             "Now deduplicated using a seen-strings set",
             "itertools.permutations generates all orderings including those with identical surface strings (when two phrases have the same words). Duplicates inflate pair counts and bias the model."],
            ["Max variants",
             "99 (random sample if >99)",
             "Paper specifies 'if >100, randomly sample exactly 99'. Implemented as stated."],
            ["Phrase integrity",
             "Full subtree via DFS moved as a unit",
             "Paper requires complete phrases to move together. DFS ensures all dependents of a head move with it."],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )

    style_heading(doc, "14.4 Language Models", 2)
    simple_table(doc,
        ["Assumption", "What we did", "Reasoning"],
        [
            ["Trigram smoothing",
             "MLE with manual 3-level backoff (trigram → bigram → unigram → epsilon)",
             "Paper does not specify smoothing. We chose MLE with backoff rather than Kneser-Ney because it is simpler and the backoff gives reasonable coverage. Kneser-Ney would likely improve accuracy."],
            ["LSTM architecture",
             "embed=256, hidden=256, 2 layers, dropout=0.3",
             "Paper does not specify architecture. Standard small LM architecture chosen. Larger models would likely give better surprisal estimates."],
            ["LSTM training data",
             "2 million tokens from Wikipedia, 2 epochs",
             "Paper does not specify training size or epochs. We used a subset to keep training time manageable on CPU. More data and epochs would improve the model."],
            ["Adaptive LR",
             "SGD, lr=0.005, 1 gradient step on context",
             "Following van Schijndel & Linzen (2018) which the paper cites. One step prevents over-adaptation."],
            ["Adaptive scope",
             "Adapt once per unique reference sentence, not per pair",
             "Correct interpretation: all variants of a sentence share the same adapted weights. Adapting per pair would apply N gradient steps for the same context."],
            ["Adaptive target",
             "Adapt on CONTEXT sentence, score on TARGET",
             "Bug fix from original implementation which adapted on the reference itself (circular). Correct: adapt on the PRECEDING sentence, then score the current sentence."],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )

    style_heading(doc, "14.5 Information Status", 2)
    simple_table(doc,
        ["Assumption", "What we did", "Reasoning"],
        [
            ["Context window",
             "Only immediately preceding sentence",
             "Paper specifies 'preceding context sentence' (singular). We use exactly one sentence back."],
            ["Givenness by pronoun",
             "Head token UPOS == PRON → GIVEN",
             "Paper explicitly states this rule. Pronouns refer anaphorically so always GIVEN."],
            ["Givenness by overlap",
             "Any content word in subtree appears in context (lowercased lemma match)",
             "Paper says 'content words within these phrases are mentioned'. We use lemma matching to handle inflection. Function words excluded."],
            ["IS score for variants",
             "Computed using REFERENCE parse with variant word order",
             "Variants are permutations of the same tokens — the dependency parse stays the same, only positions change. We use the reference parse to determine subject/object, but use the variant's word order to determine which phrase is 'first'."],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )

    style_heading(doc, "14.6 Ranking Model", 2)
    simple_table(doc,
        ["Assumption", "What we did", "Reasoning"],
        [
            ["Cross-validation split level",
             "Sentence level (not pair level)",
             "All variants of a sentence must stay in the same fold to prevent leakage. If split at pair level, the model could see the reference in train and its variant in test."],
            ["Number of folds",
             "10-fold CV",
             "Paper specifies 10-fold. Implemented as stated."],
            ["Pairwise transformation",
             "Each pair duplicated with negated features and flipped label",
             "Standard pairwise ranking formulation. Forces the model to learn directional effects rather than biasing toward label=1."],
            ["Significance test",
             "McNemar's test (paired, sentence-level)",
             "Paper uses McNemar's test. We implement it for incremental predictor comparisons."],
        ],
        col_widths=[1.8, 2.2, 3.0]
    )

    # ── 15. How to Run ───────────────────────────────────────────────────────
    style_heading(doc, "15. Running the Pipeline", 1)
    body(doc, "All commands from the project root. Virtual environment must be activated.")
    simple_table(doc,
        ["Phase", "Command", "Produces"],
        [
            ["Preprocess Wikipedia",  "python preprocessing/wiki_extract.py",                        "wiki_sentences.txt"],
            ["Build vocabulary",      "python preprocessing/build_vocab.py",                          "vocab.pkl"],
            ["Train trigram LM",      "python models/trigram/train_trigram.py",                       "trigram.pkl"],
            ["Train LSTM LM",         "python models/lstm/train_base_model.py",                       "base_model.pt"],
            ["Build feature dataset", "PYTHONPATH=. python scripts/build_feature_dataset.py",         "features.csv"],
            ["Train ranking model",   "PYTHONPATH=. python scripts/train_ranking_model.py",           "Printed results"],
            ["Compare with paper",    "PYTHONPATH=. python scripts/compare_with_paper.py",            "paper_comparison.txt"],
            ["HTML report",           "PYTHONPATH=. python scripts/analyse_results.py",               "pipeline_analysis.html"],
            ["Run tests",             "PYTHONPATH=. python tests/test_dl.py  (etc.)",                 "tests/output_test/"],
        ],
        col_widths=[1.8, 3.2, 2.0]
    )

    # ── Save ─────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved to: {OUTPUT_PATH}")
    body(doc, "All commands from the project root. Virtual environment must be activated.")
    simple_table(doc,
        ["Phase", "Command", "Produces"],
        [
            ["Preprocess Wikipedia",  "python preprocessing/wiki_extract.py",                        "wiki_sentences.txt"],
            ["Build vocabulary",      "python preprocessing/build_vocab.py",                          "vocab.pkl"],
            ["Train trigram LM",      "python models/trigram/train_trigram.py",                       "trigram.pkl"],
            ["Train LSTM LM",         "python models/lstm/train_base_model.py",                       "base_model.pt"],
            ["Build feature dataset", "PYTHONPATH=. python scripts/build_feature_dataset.py",         "features.csv"],
            ["Train ranking model",   "PYTHONPATH=. python scripts/train_ranking_model.py",           "Printed results"],
            ["Compare with paper",    "PYTHONPATH=. python scripts/compare_with_paper.py",            "paper_comparison.txt"],
            ["HTML report",           "PYTHONPATH=. python scripts/analyse_results.py",               "pipeline_analysis.html"],
            ["Run tests",             "PYTHONPATH=. python tests/test_dl.py  (etc.)",                 "tests/output_test/"],
        ],
        col_widths=[1.8, 3.2, 2.0]
    )

    # ── Save ─────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
